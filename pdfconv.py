import tkinter as tk
from tkinter import filedialog
import os
import io
import docx
from PyPDF2 import PdfReader

def pdf_to_docx(pdf_path, docx_path):
    # Open the PDF file
    with open(pdf_path, 'rb') as pdf_file:
        # Create a PDF reader object
        pdf_reader = PdfReader(pdf_file)

        # Create a new Word document
        doc = docx.Document()

        # Loop over each page in the PDF file
        for page in pdf_reader.pages:
            # Extract the text from the page
            text = page.extract_text()

            # Add the text to the Word document
            doc.add_paragraph(text)

        # Save the Word document
        doc.save(docx_path)

def browse_pdf_file():
    # Ask the user to select a PDF file
    pdf_path = filedialog.askopenfilename(filetypes=[("PDF Files", "*.pdf")])
    pdf_path_entry.delete(0, tk.END)
    pdf_path_entry.insert(0, pdf_path)

def browse_docx_file():
    # Ask the user to select a Word file
    docx_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Files", "*.docx")])
    docx_path_entry.delete(0, tk.END)
    docx_path_entry.insert(0, docx_path)

# Create a new Tkinter window
window = tk.Tk()
window.title("PDF to DOCX Converter")

# Create a label and text entry for the PDF file path
pdf_path_label = tk.Label(window, text="PDF File:")
pdf_path_label.grid(row=0, column=0)
pdf_path_entry = tk.Entry(window)
pdf_path_entry.grid(row=0, column=1)
pdf_browse_button = tk.Button(window, text="Browse", command=browse_pdf_file)
pdf_browse_button.grid(row=0, column=2)

# Create a label and text entry for the DOCX file path
docx_path_label = tk.Label(window, text="DOCX File:")
docx_path_label.grid(row=1, column=0)
docx_path_entry = tk.Entry(window)
docx_path_entry.grid(row=1, column=1)
docx_browse_button = tk.Button(window, text="Browse", command=browse_docx_file)
docx_browse_button.grid(row=1, column=2)

# Create a button to convert the PDF to DOCX
convert_button = tk.Button(window, text="Convert", command=lambda: pdf_to_docx(pdf_path_entry.get(), docx_path_entry.get()))
convert_button.grid(row=2, column=1)

def pdf_to_docx(pdf_path, docx_path):
    # Open the PDF file
    with open(pdf_path, 'rb') as pdf_file:
        # Create a PDF reader object
        pdf_reader = PdfReader(pdf_file)

        # Create a new Word document
        doc = docx.Document()

        # Loop over each page in the PDF file
        for page in pdf_reader.pages:
            # Extract the text from the page
            text = page.extract_text()

            # Add the text to the Word document
            doc.add_paragraph(text)

        # Save the Word document
        doc.save(docx_path)

    # Close the window and exit the program
    window.quit()


# Start the Tkinter event loop
window.mainloop()

